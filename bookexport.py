"""
bookexport.py
=============
Compile the conversation into a finished book:
  - build_html(messages, title) -> a single self-contained HTML string with math
    rendered by MathJax and figures embedded as images. Open it and use the
    browser's "Print -> Save as PDF" for a perfect PDF.
  - build_docx(messages, title, path) -> a Word .docx with figures and equations
    embedded as images (so math looks right in Word).

Self-contained: it re-implements the small bits it needs so it does not import
the Streamlit app.
"""

import io
import os
import re
import base64
import shutil
import subprocess
import contextlib
from pathlib import Path

_CODE_SPLIT = re.compile(r"```([\w+\-]*)\s*\n(.*?)```", re.DOTALL)


def _fig_caption(code):
    for line in code.splitlines()[:4]:
        m = re.match(r"\s*#\s*FIGURE:\s*(.+)", line, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None


def _segments(text):
    segs, pos = [], 0
    for m in _CODE_SPLIT.finditer(text or ""):
        if m.start() > pos:
            segs.append(("text", text[pos:m.start()]))
        segs.append(("code", (m.group(1) or "").lower().strip(), m.group(2)))
        pos = m.end()
    if pos < len(text or ""):
        segs.append(("text", text[pos:]))
    return segs


def render_code_to_png(code):
    """Run figure code (matplotlib/numpy/sympy/physlib) and return PNG bytes or None."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np
        try:
            import sympy as sp
        except Exception:
            sp = None
        try:
            import physlib as phys
        except Exception:
            phys = None
        plt.close("all")
        ns = {"plt": plt, "matplotlib": matplotlib, "np": np, "numpy": np,
              "sp": sp, "sympy": sp, "phys": phys, "__name__": "__main__"}
        with contextlib.redirect_stdout(io.StringIO()):
            exec(code, ns)  # noqa: S102
        nums = plt.get_fignums()
        if not nums:
            return None
        buf = io.BytesIO()
        plt.figure(nums[-1]).savefig(buf, format="png", dpi=300, bbox_inches="tight")
        plt.close("all")
        return buf.getvalue()
    except Exception:
        return None


def render_code_to_pdf(code, path):
    """Run figure code and save it as a VECTOR pdf (sharp at any size). True/False."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as np
        try:
            import sympy as sp
        except Exception:
            sp = None
        try:
            import physlib as phys
        except Exception:
            phys = None
        plt.close("all")
        ns = {"plt": plt, "matplotlib": matplotlib, "np": np, "numpy": np,
              "sp": sp, "sympy": sp, "phys": phys, "__name__": "__main__"}
        with contextlib.redirect_stdout(io.StringIO()):
            exec(code, ns)  # noqa: S102
        nums = plt.get_fignums()
        if not nums:
            return False
        plt.figure(nums[-1]).savefig(path, format="pdf", bbox_inches="tight")
        plt.close("all")
        return True
    except Exception:
        return False


def _math_to_png(latex, display=False, fontsize=15):
    """Render a LaTeX math string to PNG via matplotlib mathtext. None on failure."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        s = latex.strip()
        if not s:
            return None
        fig = plt.figure(figsize=(0.01, 0.01))
        t = fig.text(0, 0, f"${s}$", fontsize=fontsize if not display else fontsize + 3)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=300, bbox_inches="tight",
                    pad_inches=0.05, transparent=True)
        plt.close(fig)
        return buf.getvalue()
    except Exception:
        return None


# ---------------------------------------------------------------------------
# HTML  (faithful: MathJax + embedded figures; print to PDF)
# ---------------------------------------------------------------------------

def _md_to_html(text):
    """Convert prose to HTML. Uses the `markdown` package if available, else a
    light fallback. Math ($...$) is left intact for MathJax."""
    try:
        import markdown
        return markdown.markdown(text, extensions=["extra", "sane_lists"])
    except Exception:
        html = []
        for line in text.split("\n"):
            s = line.rstrip()
            if s.startswith("### "):
                html.append(f"<h3>{s[4:]}</h3>")
            elif s.startswith("## "):
                html.append(f"<h2>{s[3:]}</h2>")
            elif s.startswith("# "):
                html.append(f"<h1>{s[2:]}</h1>")
            elif s.strip() == "":
                html.append("<br>")
            else:
                s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
                html.append(f"<p>{s}</p>")
        return "\n".join(html)


def build_html(messages, title="My Book"):
    body = []
    for m in messages:
        if m.get("role") != "assistant":
            continue
        for seg in _segments(m["content"]):
            if seg[0] == "text":
                if seg[1].strip():
                    body.append(_md_to_html(seg[1]))
            else:
                _, lang, code = seg
                if lang in ("", "python", "py"):
                    png = render_code_to_png(code)
                    if png:
                        b64 = base64.b64encode(png).decode()
                        body.append(f'<p style="text-align:center">'
                                    f'<img src="data:image/png;base64,{b64}" '
                                    f'style="max-width:90%"></p>')
    mathjax = """
<script>
window.MathJax = { tex: { inlineMath: [['$','$']], displayMath: [['$$','$$']] } };
</script>
<script async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
"""
    return f"""<!doctype html><html><head><meta charset="utf-8">
<title>{title}</title>{mathjax}
<style>
body{{max-width:760px;margin:40px auto;padding:0 20px;font-family:Georgia,serif;
line-height:1.6;color:#1a1a1a}}
h1{{font-size:2em;margin-top:1.4em}} h2{{margin-top:1.3em}} img{{margin:12px 0}}
@media print{{body{{margin:0}}}}
</style></head><body>
<h1 style="text-align:center">{title}</h1>
{''.join(body)}
</body></html>"""


# ---------------------------------------------------------------------------
# DOCX  (figures + math as inline images so equations look right in Word)
# ---------------------------------------------------------------------------

_MATH_SPLIT = re.compile(r"(\$\$.+?\$\$|\$[^$]+?\$)", re.DOTALL)


def _add_rich_paragraph(doc, line):
    """Add a paragraph, rendering any $...$ math as inline images."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    parts = _MATH_SPLIT.split(line)
    for part in parts:
        if not part:
            continue
        if part.startswith("$"):
            latex = part.strip("$")
            png = _math_to_png(latex)
            if png:
                run = p.add_run()
                try:
                    run.add_picture(io.BytesIO(png), height=Pt(13))
                    continue
                except Exception:
                    pass
            p.add_run(latex)  # fallback: show the LaTeX text
        else:
            # minimal bold handling
            bold_parts = re.split(r"(\*\*.+?\*\*)", part)
            for bp in bold_parts:
                if bp.startswith("**") and bp.endswith("**"):
                    p.add_run(bp[2:-2]).bold = True
                elif bp:
                    p.add_run(bp)
    return p


def build_docx(messages, title, path):
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_heading(title, 0)
    for m in messages:
        if m.get("role") != "assistant":
            continue
        for seg in _segments(m["content"]):
            if seg[0] == "text":
                for raw in seg[1].split("\n"):
                    line = raw.rstrip()
                    if not line.strip():
                        continue
                    if line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.lstrip().startswith(("- ", "* ")):
                        doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
                    else:
                        _add_rich_paragraph(doc, line)
            else:
                _, lang, code = seg
                if lang in ("", "python", "py"):
                    png = render_code_to_png(code)
                    if png:
                        doc.add_picture(io.BytesIO(png), width=Inches(5.0))
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# LATEX  (highest quality: native LaTeX math + vector PDF figures -> PDF)
# ---------------------------------------------------------------------------

_LATEX_PREAMBLE = r"""\documentclass[11pt]{article}
\usepackage[utf8]{inputenc}
\usepackage{amsmath,amssymb,amsfonts}
\usepackage{graphicx}
\usepackage{float}
\usepackage[margin=1in]{geometry}
\usepackage{hyperref}
\title{%s}
\author{}
\date{}
\begin{document}
\maketitle
"""


_DISPLAY_MATH = re.compile(r"(\$\$.+?\$\$|\\\[.+?\\\])", re.DOTALL)
_EMPH_LINE = re.compile(r"^\s*\*{1,2}([^*].*?)\*{1,2}\s*$")


def _latex_escape_text(s):
    """Escape LaTeX specials in prose, leaving inline $...$ math untouched and
    converting **bold** / *italic* markdown."""
    parts = re.split(r"(\$[^$]+?\$)", s)
    out = []
    for p in parts:
        if not p:
            continue
        if p.startswith("$") and p.endswith("$"):
            out.append(p)            # inline math: already LaTeX, pass through
            continue
        for ch, rep in [("&", r"\&"), ("%", r"\%"), ("#", r"\#"), ("_", r"\_")]:
            p = p.replace(ch, rep)
        p = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", p)   # bold first
        p = re.sub(r"\*([^*\n]+?)\*", r"\\textit{\1}", p)  # then italic
        out.append(p)
    return "".join(out)


def build_latex(messages, title, out_dir):
    """
    Write a LaTeX book to out_dir: book.tex + figures/*.pdf (vector). If a
    LaTeX engine (pdflatex) is available, also compile out_dir/book.pdf.
    Returns dict {dir, tex, pdf_or_None, n_figures}.
    """
    out = Path(out_dir)
    figs = out / "figures"
    figs.mkdir(parents=True, exist_ok=True)
    body = []
    in_list = False
    fig_n = 0

    def _close_list():
        nonlocal in_list
        if in_list:
            body.append(r"\end{itemize}")
            in_list = False

    for m in messages:
        if m.get("role") != "assistant":
            continue
        for seg in _segments(m["content"]):
            if seg[0] == "text":
                # Protect multi-line display math ($$...$$ / \[...\]) before any
                # line processing, so subscripts etc. are never escaped.
                for chunk in _DISPLAY_MATH.split(seg[1]):
                    if not chunk or not chunk.strip():
                        continue
                    cs = chunk.lstrip()
                    if cs.startswith("$$") or cs.startswith("\\["):
                        _close_list()
                        body.append(chunk.strip())          # raw display math
                        continue
                    for raw in chunk.split("\n"):
                        line = raw.rstrip()
                        if not line.strip():
                            _close_list(); body.append(""); continue
                        if line.strip() in ("---", "***", "___", "—", "--"):
                            _close_list(); body.append(r"\medskip"); continue
                        if line.startswith("### "):
                            _close_list(); body.append(r"\subsubsection*{%s}" % _latex_escape_text(line[4:]))
                        elif line.startswith("## "):
                            _close_list(); body.append(r"\subsection*{%s}" % _latex_escape_text(line[3:]))
                        elif line.startswith("# "):
                            _close_list(); body.append(r"\section*{%s}" % _latex_escape_text(line[2:]))
                        elif line.lstrip().startswith(("- ", "+ ")):
                            if not in_list:
                                body.append(r"\begin{itemize}"); in_list = True
                            body.append(r"\item " + _latex_escape_text(line.lstrip()[2:]))
                        elif _EMPH_LINE.match(line):
                            # a whole line wrapped in * or ** = a heading the model meant
                            _close_list()
                            body.append(r"\subsection*{%s}" % _latex_escape_text(_EMPH_LINE.match(line).group(1)))
                        else:
                            _close_list()
                            body.append(_latex_escape_text(line))
            else:
                _close_list()
                _, lang, code = seg
                if lang not in ("", "python", "py"):
                    continue
                fig_n += 1
                fname = f"fig{fig_n}.pdf"
                if render_code_to_pdf(code, str(figs / fname)):
                    cap = _fig_caption(code)
                    body.append(r"\begin{figure}[H]\centering")
                    body.append(r"\includegraphics[width=0.78\textwidth]{figures/%s}" % fname)
                    if cap:
                        body.append(r"\caption{%s}" % _latex_escape_text(cap))
                    body.append(r"\end{figure}")
    _close_list()

    tex = (_LATEX_PREAMBLE % _latex_escape_text(title)) + "\n".join(body) + "\n\\end{document}\n"
    tex_path = out / "book.tex"
    tex_path.write_text(tex, encoding="utf-8")

    pdf_path = None
    if shutil.which("pdflatex"):
        try:
            for _ in range(2):  # twice for references/toc
                subprocess.run(["pdflatex", "-interaction=nonstopmode", "book.tex"],
                               cwd=str(out), stdout=subprocess.DEVNULL,
                               stderr=subprocess.DEVNULL, timeout=120)
            if (out / "book.pdf").exists():
                pdf_path = str(out / "book.pdf")
        except Exception:
            pdf_path = None
    return {"dir": str(out), "tex": str(tex_path), "pdf": pdf_path, "n_figures": fig_n}
